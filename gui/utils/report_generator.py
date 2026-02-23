"""
ReportGenerator — Word (.docx) formatida rasmiy hisobot yaratish.
Belgilangan sana oralig'ida pereezd monitoring statistikasini eksport qiladi.
Professional dizayn: rangli jadvallar, chiroyli sarlavhalar, zebra qatorlar.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime, date, timedelta

# ─── Ranglar ──────────────────────────────────────────────
BRAND = RGBColor(0x0D, 0x47, 0x6E)       # Quyuq ko'k
BRAND_LIGHT = RGBColor(0x15, 0x6B, 0x9A)  # O'rtacha ko'k
ACCENT_GREEN = RGBColor(0x27, 0xAE, 0x60)
ACCENT_ORANGE = RGBColor(0xE6, 0x7E, 0x22)
ACCENT_RED = RGBColor(0xC0, 0x39, 0x2B)
TEXT_DARK = RGBColor(0x2C, 0x3E, 0x50)
TEXT_GRAY = RGBColor(0x7F, 0x8C, 0x8D)
TEXT_LIGHT = RGBColor(0x95, 0xA5, 0xA6)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Hex ranglar (shading uchun)
HDR_BG = "0D476E"     # Jadval header fon
HDR_BG2 = "156B9A"    # Ikkinchi header variant
ROW_EVEN = "F0F7FB"   # Juft qator fon
ROW_ODD = "FFFFFF"    # Toq qator fon
LIGHT_BLUE_BG = "EBF5FB"
LIGHT_GREEN_BG = "EAFAF1"
LIGHT_ORANGE_BG = "FEF5E7"
ACCENT_LINE = "0D476E"
GRAY_LINE = "BDC3C7"


def generate_report(config_manager, stats_db, date_from: str, date_to: str,
                    file_path: str) -> bool:
    """Word hisobot yaratish."""
    try:
        doc = Document()

        # Default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)
        font.color.rgb = TEXT_DARK

        # Sahifa sozlamalari
        section = doc.sections[0]
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

        crossings = config_manager.get_crossings()

        _add_title_page(doc, date_from, date_to, crossings)
        _add_summary(doc, stats_db, crossings, date_from, date_to)

        for i, crossing in enumerate(crossings):
            _add_crossing_section(doc, stats_db, crossing, date_from, date_to,
                                  section_num=i + 2)

        _add_footer(doc)

        doc.save(file_path)
        return True

    except Exception as e:
        print(f"[ReportGenerator] Error: {e}")
        import traceback
        traceback.print_exc()
        return False


# ═══════════════════════════════════════════════════════════
#  TITLE PAGE
# ═══════════════════════════════════════════════════════════

def _add_title_page(doc, date_from, date_to, crossings):
    for _ in range(3):
        doc.add_paragraph("")

    # Dekorativ chiziq (tepada)
    _add_colored_line(doc, ACCENT_LINE, width=4)

    doc.add_paragraph("")

    # Tashkilot nomi
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("O'ZBEKISTON TEMIR YO'LLARI")
    run.font.size = Pt(13)
    run.font.color.rgb = TEXT_GRAY
    run.font.name = 'Calibri'
    _set_spacing(p, before=0, after=4)

    # Asosiy sarlavha
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TEMIR YO'L PEREEZD")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = BRAND
    run.font.name = 'Calibri'
    _set_spacing(title, before=0, after=0)

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title2.add_run("MONITORING HISOBOTI")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = BRAND
    run.font.name = 'Calibri'
    _set_spacing(title2, before=0, after=12)

    # Yana chiziq
    _add_colored_line(doc, ACCENT_LINE, width=2)

    doc.add_paragraph("")

    # Hisobot davri — rangli quti
    period = doc.add_paragraph()
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = period.add_run(f"Hisobot davri:  {_format_date(date_from)}  —  {_format_date(date_to)}")
    run.font.size = Pt(14)
    run.font.color.rgb = BRAND_LIGHT
    run.font.name = 'Calibri'
    _set_spacing(period, before=0, after=16)

    # Ma'lumot jadvali (2x2 karta ko'rinishi)
    total_cams = sum(len(cr.get("cameras", [])) for cr in crossings)
    info_table = doc.add_table(rows=1, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(info_table)

    # Pereezdlar soni
    c1 = info_table.rows[0].cells[0]
    _add_info_card(c1, str(len(crossings)), "Pereezdlar soni", LIGHT_BLUE_BG)

    # Kameralar soni
    c2 = info_table.rows[0].cells[1]
    _add_info_card(c2, str(total_cams), "Kameralar soni", LIGHT_GREEN_BG)

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Yaratilgan vaqt
    gen = doc.add_paragraph()
    gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = gen.add_run(f"Hisobot yaratilgan: {datetime.now().strftime('%d.%m.%Y  %H:%M')}")
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT_LIGHT
    run.italic = True
    _set_spacing(gen, before=0, after=0)

    sys_p = doc.add_paragraph()
    sys_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sys_p.add_run("RailSafe Monitoring System v2.0")
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT_LIGHT
    run.italic = True

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════
#  SUMMARY
# ═══════════════════════════════════════════════════════════

def _add_summary(doc, stats_db, crossings, date_from, date_to):
    _section_header(doc, "1", "UMUMIY STATISTIKA")

    grand_light = 0
    grand_heavy = 0
    for cr in crossings:
        l, h = stats_db.get_date_range_total(cr["id"], date_from, date_to)
        grand_light += l
        grand_heavy += h
    grand_total = grand_light + grand_heavy

    # 3 ta karta: Jami, Yengil, Og'ir
    cards = doc.add_table(rows=1, cols=3)
    cards.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(cards)

    _add_stat_card(cards.rows[0].cells[0], str(grand_total),
                   "JAMI TRANSPORT", LIGHT_BLUE_BG, BRAND)
    _add_stat_card(cards.rows[0].cells[1], str(grand_light),
                   "YENGIL TRANSPORT", LIGHT_GREEN_BG, ACCENT_GREEN)
    _add_stat_card(cards.rows[0].cells[2], str(grand_heavy),
                   "OG'IR TRANSPORT", LIGHT_ORANGE_BG, ACCENT_ORANGE)

    _add_spacer(doc, 8)

    # Pereezdlar bo'yicha jadval
    _sub_header(doc, "Pereezdlar bo'yicha taqsimot")

    headers = ["#", "Pereezd nomi", "Joylashuv", "Yengil", "Og'ir", "Jami", "Kameralar"]
    rows_data = []
    for idx, cr in enumerate(crossings):
        l, h = stats_db.get_date_range_total(cr["id"], date_from, date_to)
        rows_data.append([
            str(idx + 1),
            cr.get("name", f"Pereezd #{cr['id']}"),
            cr.get("location", "—"),
            str(l),
            str(h),
            str(l + h),
            str(len(cr.get("cameras", [])))
        ])

    _add_styled_table(doc, headers, rows_data,
                      col_widths=[0.6, 2.5, 2.0, 1.5, 1.5, 1.5, 1.2],
                      right_align_cols=[3, 4, 5, 6],
                      bold_col=5)

    _add_spacer(doc, 4)


# ═══════════════════════════════════════════════════════════
#  CROSSING SECTION
# ═══════════════════════════════════════════════════════════

def _add_crossing_section(doc, stats_db, crossing, date_from, date_to,
                          section_num=2):
    cid = crossing["id"]
    name = crossing.get("name", f"Pereezd #{cid}")
    location = crossing.get("location", "")
    cameras = crossing.get("cameras", [])

    doc.add_page_break()
    _section_header(doc, str(section_num), f"PEREEZD: {name.upper()}")

    if location:
        p = doc.add_paragraph()
        run = p.add_run("Joylashuv: ")
        run.font.size = Pt(10)
        run.font.color.rgb = TEXT_GRAY
        run = p.add_run(location)
        run.font.size = Pt(10)
        run.font.color.rgb = TEXT_DARK
        run.bold = True
        _set_spacing(p, before=0, after=6)

    # Jami statistika kartalari
    light, heavy = stats_db.get_date_range_total(cid, date_from, date_to)
    total = light + heavy

    cards = doc.add_table(rows=1, cols=3)
    cards.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(cards)

    _add_stat_card(cards.rows[0].cells[0], str(total),
                   "JAMI", LIGHT_BLUE_BG, BRAND)
    _add_stat_card(cards.rows[0].cells[1], str(light),
                   "YENGIL", LIGHT_GREEN_BG, ACCENT_GREEN)
    _add_stat_card(cards.rows[0].cells[2], str(heavy),
                   "OG'IR", LIGHT_ORANGE_BG, ACCENT_ORANGE)

    _add_spacer(doc, 8)

    # Kameralar statistikasi
    if cameras:
        _sub_header(doc, "Kameralar statistikasi")
        cam_headers = ["#", "Kamera", "Turi", "Yengil", "Og'ir", "Jami"]
        cam_rows = []
        for idx, cam in enumerate(cameras):
            cn = cam.get("name", "?")
            ct = "Asosiy" if cam.get("type") == "main" else "Qo'shimcha"
            cl, ch = stats_db.get_date_range_camera(cid, cn, date_from, date_to)
            cam_rows.append([str(idx + 1), cn, ct, str(cl), str(ch), str(cl + ch)])

        _add_styled_table(doc, cam_headers, cam_rows,
                          col_widths=[0.6, 1.8, 2.0, 1.8, 1.8, 1.8],
                          right_align_cols=[3, 4, 5],
                          bold_col=5)
        _add_spacer(doc, 6)

    # Kunlik statistika
    daily_data = stats_db.get_date_range_daily(cid, date_from, date_to)
    if daily_data:
        _sub_header(doc, "Kunlik statistika")

        daily_headers = ["#", "Sana", "Hafta kuni", "Yengil", "Og'ir", "Jami"]
        days_uz = ["Dushanba", "Seshanba", "Chorshanba", "Payshanba",
                   "Juma", "Shanba", "Yakshanba"]
        daily_rows = []
        for idx, d in enumerate(daily_data):
            try:
                dt = datetime.strptime(d["date"], "%Y-%m-%d")
                day_name = days_uz[dt.weekday()]
            except (ValueError, TypeError):
                day_name = "—"
            daily_rows.append([
                str(idx + 1),
                _format_date(d["date"]),
                day_name,
                str(d["light"]),
                str(d["heavy"]),
                str(d["light"] + d["heavy"])
            ])

        _add_styled_table(doc, daily_headers, daily_rows,
                          col_widths=[0.6, 1.8, 2.2, 1.5, 1.5, 1.5],
                          right_align_cols=[3, 4, 5],
                          bold_col=5)


# ═══════════════════════════════════════════════════════════
#  FOOTER
# ═══════════════════════════════════════════════════════════

def _add_footer(doc):
    _add_spacer(doc, 12)
    _add_colored_line(doc, GRAY_LINE, width=1)
    _add_spacer(doc, 4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— Hisobot tugadi —")
    run.font.size = Pt(10)
    run.font.color.rgb = TEXT_LIGHT
    run.italic = True
    _set_spacing(p, before=0, after=4)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run("RailSafe Monitoring System  |  ")
    run.font.size = Pt(8)
    run.font.color.rgb = TEXT_LIGHT
    run = p2.add_run(f"Yaratilgan: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    run.font.size = Pt(8)
    run.font.color.rgb = TEXT_LIGHT


# ═══════════════════════════════════════════════════════════
#  STYLED TABLE
# ═══════════════════════════════════════════════════════════

def _add_styled_table(doc, headers, rows_data, col_widths=None,
                      right_align_cols=None, bold_col=None):
    """Professional jadval: rangli header, zebra qatorlar, to'g'ri alignment."""
    right_align_cols = right_align_cols or []
    n_cols = len(headers)
    n_rows = len(rows_data) + 1  # +1 header uchun

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Ustun kengliklarini sozlash
    if col_widths:
        for i, w in enumerate(col_widths):
            if i < n_cols:
                for row in table.rows:
                    row.cells[i].width = Cm(w)

    # ─── Header qatori ───
    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_shading(cell, HDR_BG)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(p, before=2, after=2)

        run = p.add_run(h_text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'

    # ─── Ma'lumot qatorlari ───
    for r_idx, row_data in enumerate(rows_data):
        bg = ROW_EVEN if r_idx % 2 == 0 else ROW_ODD
        for c_idx in range(n_cols):
            cell = table.rows[r_idx + 1].cells[c_idx]
            _set_cell_shading(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p = cell.paragraphs[0]
            if c_idx in right_align_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_spacing(p, before=1, after=1)

            value = row_data[c_idx] if c_idx < len(row_data) else ""
            run = p.add_run(value)
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            run.font.color.rgb = TEXT_DARK

            if bold_col is not None and c_idx == bold_col:
                run.bold = True
                run.font.color.rgb = BRAND

    # Jadval chegaralari
    _set_table_borders(table)


# ═══════════════════════════════════════════════════════════
#  HELPERS — Components
# ═══════════════════════════════════════════════════════════

def _section_header(doc, num, text):
    """Rangli raqam + sarlavha"""
    p = doc.add_paragraph()
    _set_spacing(p, before=0, after=8)

    # Raqam
    run = p.add_run(f"  {num}  ")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = WHITE
    run.font.name = 'Calibri'
    # Raqam uchun shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HDR_BG}" w:val="clear"/>')
    run._element.get_or_add_rPr().append(shading)

    # Sarlavha matni
    run = p.add_run(f"  {text}")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = BRAND
    run.font.name = 'Calibri'

    # Tagiga chiziq
    _add_colored_line(doc, ACCENT_LINE, width=2)
    _add_spacer(doc, 6)


def _sub_header(doc, text):
    """Kichik sarlavha — chap chiziq bilan"""
    p = doc.add_paragraph()
    _set_spacing(p, before=4, after=4)
    # Chapda vertikal chiziq effekti
    run = p.add_run("▌ ")
    run.font.color.rgb = BRAND_LIGHT
    run.font.size = Pt(11)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_DARK
    run.font.name = 'Calibri'


def _add_info_card(cell, value, label, bg_color):
    """Sarlavha sahifasidagi info karta (cell ichida)"""
    _set_cell_shading(cell, bg_color)
    _set_cell_padding(cell, top=200, bottom=200, left=150, right=150)

    # Qiymat
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(value)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = BRAND
    run.font.name = 'Calibri'

    # Label
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run(label)
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT_GRAY
    run.font.name = 'Calibri'


def _add_stat_card(cell, value, label, bg_color, value_color):
    """Statistika karta (kichikroq, 3 ustunli)"""
    _set_cell_shading(cell, bg_color)
    _set_cell_padding(cell, top=150, bottom=150, left=100, right=100)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, before=0, after=0)
    run = p.add_run(value)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = value_color
    run.font.name = 'Calibri'

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p2, before=0, after=0)
    run = p2.add_run(label)
    run.font.size = Pt(8)
    run.font.color.rgb = TEXT_GRAY
    run.font.name = 'Calibri'
    run.bold = True


# ═══════════════════════════════════════════════════════════
#  HELPERS — Low level
# ═══════════════════════════════════════════════════════════

def _format_date(date_str):
    """2026-02-01 → 01.02.2026"""
    try:
        d = datetime.strptime(date_str, "%Y-%m-%d")
        return d.strftime("%d.%m.%Y")
    except (ValueError, TypeError):
        return date_str


def _set_cell_shading(cell, color_hex):
    """Yacheyka fon rangini o'rnatish"""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_padding(cell, top=0, bottom=0, left=0, right=0):
    """Yacheyka ichki bo'shligi (twips)"""
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>')
    tc_pr.append(margins)


def _set_table_borders(table):
    """Jadval chegaralarini ingichka qilish"""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{GRAY_LINE}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{GRAY_LINE}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{GRAY_LINE}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{GRAY_LINE}"/>'
        f'  <w:insideH w:val="single" w:sz="2" w:space="0" w:color="{GRAY_LINE}"/>'
        f'  <w:insideV w:val="single" w:sz="2" w:space="0" w:color="{GRAY_LINE}"/>'
        f'</w:tblBorders>')
    tbl_pr.append(borders)


def _remove_table_borders(table):
    """Jadval chegaralarini olib tashlash"""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>')
    tbl_pr.append(borders)


def _add_colored_line(doc, color_hex, width=2):
    """Gorizontal rangli chiziq"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, before=0, after=0)
    # Pastki chegara orqali chiziq effekti
    pPr = p._p.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{width * 4}" w:space="1" w:color="{color_hex}"/>'
        f'</w:pBdr>')
    pPr.append(borders)


def _add_spacer(doc, pt_size):
    """Bo'sh joy qo'shish"""
    p = doc.add_paragraph("")
    _set_spacing(p, before=0, after=pt_size)


def _set_spacing(paragraph, before=0, after=0):
    """Paragraf bo'shlig'ini sozlash (pt)"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
