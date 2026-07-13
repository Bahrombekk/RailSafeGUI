"""
ReportGenerator — Word (.docx) formatida rasmiy hisobot yaratish.
Belgilangan sana oralig'ida pereezd monitoring statistikasini eksport qiladi.
Professional dizayn: rangli banner, accent kartalar, vizual grafiklar.
"""

import logging

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime, date, timedelta

from app.utils.language import t

logger = logging.getLogger("RailSafe.reports")

# ─── Ranglar (HTML dizaynga mos) ──────────────────────────
# Header: #0d2b4e / #1a3d6b / #1558a0  (HTML bilan bir xil)
BRAND       = RGBColor(0x1E, 0x3A, 0x5F)   # HTML: #1e3a5f karta sarlavha
BRAND_LIGHT = RGBColor(0x1D, 0x4E, 0xD8)   # HTML: #1d4ed8 ko'k accent
BRAND_PALE  = RGBColor(0xA8, 0xC8, 0xE8)   # HTML: #a8c8e8 header pale
ACCENT_GREEN  = RGBColor(0x10, 0xB9, 0x81)  # HTML: #10b981
ACCENT_ORANGE = RGBColor(0xF5, 0x9E, 0x0B)  # HTML: #f59e0b
ACCENT_RED    = RGBColor(0xEF, 0x44, 0x44)  # HTML: #ef4444
ACCENT_TEAL   = RGBColor(0x7C, 0x3A, 0xED)  # HTML: #7c3aed (poyezdlar uchun)
TEXT_DARK  = RGBColor(0x11, 0x18, 0x27)    # HTML: #111827 (card qiymati)
TEXT_MID   = RGBColor(0x37, 0x41, 0x51)    # HTML: #374151
TEXT_GRAY  = RGBColor(0x6B, 0x72, 0x80)    # HTML: #6b7280
TEXT_LIGHT = RGBColor(0x9C, 0xA3, 0xAF)    # HTML: #9ca3af
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

# Hex ranglar (banner/border uchun) — HTML bilan bir xil
HDR_BG        = "0D2B4E"   # HTML: #0d2b4e  (asosiy banner)
HDR_BG2       = "1A3D6B"   # HTML: #1a3d6b  (top bar)
HDR_BG3       = "1558A0"   # HTML: #1558a0  (bottom bar)
HDR_TEAL      = "0D9488"   # HTML: #0d9488
HDR_GREEN     = "10B981"   # HTML: #10b981
HDR_ORANGE    = "F59E0B"   # HTML: #f59e0b
ROW_EVEN      = "F9FAFB"   # HTML: #f9fafb (tr:nth-child(even))
ROW_ODD       = "FFFFFF"
TABLE_HDR_BG  = "F1F5F9"   # HTML: #f1f5f9 (thead)
SECTION_BG    = "FFFFFF"   # HTML: .section white
CARD_BG       = "FFFFFF"   # HTML: .card white
LIGHT_BG      = "F8FAFC"   # HTML: .mini-stat #f8fafc
GRAY_LINE     = "E5E7EB"   # HTML: #e5e7eb border

# Eski nomlar — CARD_BG ga yo'naltirilgan (orqaga mos)
LIGHT_BLUE_BG   = CARD_BG
LIGHT_GREEN_BG  = CARD_BG
LIGHT_ORANGE_BG = CARD_BG
LIGHT_TEAL_BG   = CARD_BG

# Accent border colors (HTML card border-top ranglar)
ACCENT_BLUE_HEX   = "1D4ED8"   # HTML: #1d4ed8
ACCENT_GREEN_HEX  = "10B981"   # HTML: #10b981
ACCENT_ORANGE_HEX = "F59E0B"   # HTML: #f59e0b
ACCENT_TEAL_HEX   = "7C3AED"   # HTML: #7c3aed


def generate_report(config_manager, stats_db, date_from: str, date_to: str,
                    file_path: str) -> bool:
    """Word hisobot yaratish."""
    try:
        doc = Document()

        # Default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)
        font.color.rgb = TEXT_DARK

        # Sahifa sozlamalari
        section = doc.sections[0]
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

        crossings = config_manager.get_crossings()

        _add_title_page(doc, date_from, date_to, crossings)
        _add_summary(doc, stats_db, crossings, date_from, date_to)

        for i, crossing in enumerate(crossings):
            _add_crossing_section(doc, stats_db, crossing, date_from, date_to,
                                  section_num=i + 2)

        _add_footer(doc)

        doc.save(file_path)
        return True

    except Exception:
        logger.exception("[ReportGenerator] Hisobot yaratishda xato")
        return False


# ═══════════════════════════════════════════════════════════
#  TITLE PAGE
# ═══════════════════════════════════════════════════════════

def _add_title_page(doc, date_from, date_to, crossings):
    total_cams = sum(len(cr.get("cameras", [])) for cr in crossings)

    # ── Ko'k banner (3 qator: org | sarlavha | davr) ────────
    banner = doc.add_table(rows=3, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(banner)
    _set_table_width(banner, Cm(17))

    # Qator 0: Tashkilot nomi (HTML: header-top-bar #1a3d6b)
    c0 = banner.rows[0].cells[0]
    _set_cell_shading(c0, HDR_BG2)
    _set_cell_padding(c0, top=180, bottom=100, left=400, right=400)
    p = c0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(t("rpt.org_name").upper())
    run.font.size = Pt(9)
    run.font.color.rgb = BRAND_PALE
    run.font.name = 'Calibri'
    _set_spacing(p, before=0, after=0)

    # Qator 1: Asosiy sarlavha (HTML: header-main #0d2b4e)
    c1 = banner.rows[1].cells[0]
    _set_cell_shading(c1, HDR_BG)
    _set_cell_padding(c1, top=100, bottom=80, left=400, right=400)

    pb = c1.paragraphs[0]
    pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pb.add_run(t("rpt.auto_system"))
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x90, 0xB8, 0xD8)
    run.font.name = 'Calibri'
    _set_spacing(pb, before=0, after=4)

    pt = c1.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pt.add_run(t("rpt.main_title").upper())
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = WHITE
    run.font.name = 'Calibri'
    _set_spacing(pt, before=0, after=2)

    ps = c1.add_paragraph()
    ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ps.add_run(t("rpt.subtitle"))
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x90, 0xB8, 0xD8)
    run.font.name = 'Calibri'
    _set_spacing(ps, before=0, after=0)

    # Qator 2: Davr (HTML: header-bottom-bar #1558a0)
    c2 = banner.rows[2].cells[0]
    _set_cell_shading(c2, HDR_BG3)
    _set_cell_padding(c2, top=120, bottom=120, left=400, right=400)
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run(
        f"{t('rpt.period')}:   {_format_date(date_from)}  —  {_format_date(date_to)}")
    run.font.size = Pt(11)
    run.font.color.rgb = WHITE
    run.font.name = 'Calibri'
    _set_spacing(p2, before=0, after=0)

    _add_spacer(doc, 14)

    # ── 4 ta info karta ──────────────────────────────────────
    info = doc.add_table(rows=1, cols=4)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(info)

    _add_banner_card(info.rows[0].cells[0],
                     str(len(crossings)), t("rpt.pereezds"), ACCENT_BLUE_HEX)
    _add_banner_card(info.rows[0].cells[1],
                     str(total_cams), t("rpt.cameras"), "156B9A")
    _add_banner_card(info.rows[0].cells[2],
                     _format_date(date_from), t("rpt.start"), ACCENT_GREEN_HEX)
    _add_banner_card(info.rows[0].cells[3],
                     _format_date(date_to), t("rpt.end"), ACCENT_TEAL_HEX)

    _add_spacer(doc, 20)

    # Yaratilgan vaqt
    gen = doc.add_paragraph()
    gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = gen.add_run(
        t("rpt.report_created_at", dt=datetime.now().strftime('%d.%m.%Y  %H:%M')))
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT_LIGHT
    run.italic = True
    _set_spacing(gen, before=0, after=2)

    sys_p = doc.add_paragraph()
    sys_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sys_p.add_run(t("rpt.system_ver"))
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT_LIGHT
    run.italic = True

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════
#  SUMMARY
# ═══════════════════════════════════════════════════════════

def _add_summary(doc, stats_db, crossings, date_from, date_to):
    _section_header(doc, "1", t("rpt.sec_summary").upper())

    grand_light  = 0
    grand_heavy  = 0
    grand_trains = 0
    for cr in crossings:
        l, h = stats_db.get_date_range_total(cr["id"], date_from, date_to)
        grand_light += l
        grand_heavy += h
        ts = stats_db.get_train_range_stats(cr["id"], date_from, date_to)
        grand_trains += ts["count"]
    grand_total = grand_light + grand_heavy

    # ── 4 ta statistika kartasi ──────────────────────────────
    cards = doc.add_table(rows=1, cols=4)
    cards.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(cards)

    _add_stat_card(cards.rows[0].cells[0], str(grand_total),
                   t("rpt.card_total_transport").upper(), LIGHT_BLUE_BG,   BRAND,
                   ACCENT_BLUE_HEX)
    _add_stat_card(cards.rows[0].cells[1], str(grand_light),
                   t("rpt.light").upper(),          LIGHT_GREEN_BG,  ACCENT_GREEN,
                   ACCENT_GREEN_HEX)
    _add_stat_card(cards.rows[0].cells[2], str(grand_heavy),
                   t("rpt.heavy").upper(),           LIGHT_ORANGE_BG, ACCENT_ORANGE,
                   ACCENT_ORANGE_HEX)
    _add_stat_card(cards.rows[0].cells[3], str(grand_trains),
                   t("rpt.card_trains").upper(),       LIGHT_TEAL_BG,   ACCENT_TEAL,
                   ACCENT_TEAL_HEX)

    _add_spacer(doc, 10)

    # ── Pereezdlar qiyosiy tahlili ───────────────────────────
    _sub_header(doc, t("rpt.compare_analysis"))

    cr_data = []
    max_total = 1
    for cr in crossings:
        l, h = stats_db.get_date_range_total(cr["id"], date_from, date_to)
        ts = stats_db.get_train_range_stats(cr["id"], date_from, date_to)
        tot = l + h
        if tot > max_total:
            max_total = tot
        cr_data.append((cr, l, h, tot, ts["count"]))

    # Jadval — # | Pereezd | Yengil | Og'ir | Jami | ████ grafik | Poyezdlar
    headers = [t("rpt.col_num"), t("rpt.col_pereezd"), t("rpt.light"),
               t("rpt.heavy"), t("rpt.total"), t("rpt.col_ratio"), t("rpt.col_trains")]
    rows_data = []
    for idx, (cr, l, h, tot, trains) in enumerate(cr_data):
        bar_len = int(tot / max_total * 20) if max_total > 0 else 0
        bar = "█" * bar_len + "░" * (20 - bar_len)
        rows_data.append([
            str(idx + 1),
            cr.get("name", f"Pereezd #{cr['id']}"),
            str(l), str(h), str(tot),
            bar,
            str(trains),
        ])

    _add_styled_table(doc, headers, rows_data,
                      col_widths=[0.5, 2.8, 1.3, 1.3, 1.3, 4.0, 1.5],
                      right_align_cols=[2, 3, 4, 6],
                      bold_col=4,
                      bar_col=5)

    _add_spacer(doc, 6)


# ═══════════════════════════════════════════════════════════
#  CROSSING SECTION
# ═══════════════════════════════════════════════════════════

def _add_crossing_section(doc, stats_db, crossing, date_from, date_to,
                          section_num=2):
    cid      = crossing["id"]
    name     = crossing.get("name", f"Pereezd #{cid}")
    location = crossing.get("location", "")
    cameras  = crossing.get("cameras", [])

    doc.add_page_break()

    # ── Pereezd banner — HTML .section + .section-header uslubi ─
    bann = doc.add_table(rows=1, cols=2)
    bann.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_table_borders(bann)
    _set_table_width(bann, Cm(17))

    # Chap: ko'k chiziq (HTML sec-title ::before)
    bl = bann.rows[0].cells[0]
    _set_cell_shading(bl, ACCENT_BLUE_HEX)
    bl.width = Cm(0.15)
    _set_cell_padding(bl, top=80, bottom=80, left=0, right=0)
    bl.paragraphs[0].add_run("")
    _set_spacing(bl.paragraphs[0], before=0, after=0)

    # O'ng: nom + joylashuv
    br = bann.rows[0].cells[1]
    _set_cell_shading(br, "FFFFFF")
    _set_cell_padding(br, top=100, bottom=80, left=180, right=180)

    bp0 = br.paragraphs[0]
    bp0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = bp0.add_run(f"{section_num}.  {name}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = BRAND        # HTML: .section-title #1e3a5f
    run.font.name = 'Calibri'
    _set_spacing(bp0, before=0, after=2)

    if location:
        bp1 = br.add_paragraph()
        bp1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = bp1.add_run(f"📡  {location}  ·  {_format_date(date_from)} — {_format_date(date_to)}")
        run.font.size = Pt(9)
        run.font.color.rgb = TEXT_GRAY   # HTML: .section-sub #6b7280
        run.font.name = 'Calibri'
        _set_spacing(bp1, before=0, after=0)

    _add_spacer(doc, 10)

    # ── Jami transport kartalari ─────────────────────────────
    light, heavy = stats_db.get_date_range_total(cid, date_from, date_to)
    total = light + heavy
    train_stats = stats_db.get_train_range_stats(cid, date_from, date_to)

    _block_label(doc, t("rpt.transport_stats").upper(), ACCENT_BLUE_HEX)

    cards = doc.add_table(rows=1, cols=4)
    cards.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(cards)

    _add_stat_card(cards.rows[0].cells[0], str(total),
                   t("rpt.total").upper(),   LIGHT_BLUE_BG,   BRAND,         ACCENT_BLUE_HEX)
    _add_stat_card(cards.rows[0].cells[1], str(light),
                   t("rpt.light").upper(), LIGHT_GREEN_BG,  ACCENT_GREEN,  ACCENT_GREEN_HEX)
    _add_stat_card(cards.rows[0].cells[2], str(heavy),
                   t("rpt.heavy").upper(),  LIGHT_ORANGE_BG, ACCENT_ORANGE, ACCENT_ORANGE_HEX)
    _add_stat_card(cards.rows[0].cells[3], str(train_stats["count"]),
                   t("rpt.card_trains").upper(), LIGHT_TEAL_BG, ACCENT_TEAL,  ACCENT_TEAL_HEX)

    _add_spacer(doc, 10)

    # ── Kameralar statistikasi ───────────────────────────────
    if cameras:
        _block_label(doc, t("rpt.cam_stats").upper(), "156B9A")
        cam_headers = [t("rpt.col_num"), t("rpt.col_camera"), t("rpt.col_type"),
                       t("rpt.light"), t("rpt.heavy"), t("rpt.total")]
        cam_rows = []
        for idx, cam in enumerate(cameras):
            cn = cam.get("name", "?")
            ct = t("crossing.type.main") if cam.get("type") == "main" else t("crossing.type.additional")
            cl, ch = stats_db.get_date_range_camera(cid, cn, date_from, date_to)
            cam_rows.append([str(idx + 1), cn, ct, str(cl), str(ch), str(cl + ch)])

        _add_styled_table(doc, cam_headers, cam_rows,
                          col_widths=[0.6, 2.8, 2.2, 1.8, 1.8, 1.8],
                          right_align_cols=[3, 4, 5],
                          bold_col=5)
        _add_spacer(doc, 10)

    # ── Kunlik statistika ────────────────────────────────────
    daily_data = stats_db.get_date_range_daily(cid, date_from, date_to)
    if daily_data:
        _block_label(doc, t("rpt.daily_stats").upper(), HDR_GREEN)

        # Hafta kunlari nomlari (0=Dushanba ... 6=Yakshanba)
        day_names = [t(f"rpt.weekday_{i}") for i in range(7)]
        max_day = max((d["light"] + d["heavy"] for d in daily_data), default=1) or 1

        daily_headers = [t("rpt.col_num"), t("rpt.col_date"), t("rpt.col_day"),
                         t("rpt.light"), t("rpt.heavy"), t("rpt.total"), t("rpt.col_chart")]
        daily_rows = []
        for idx, d in enumerate(daily_data):
            try:
                dt = datetime.strptime(d["date"], "%Y-%m-%d")
                day_name = day_names[dt.weekday()]
            except (ValueError, TypeError):
                day_name = "—"
            tot = d["light"] + d["heavy"]
            bar_len = int(tot / max_day * 16) if max_day > 0 else 0
            bar = "█" * bar_len
            daily_rows.append([
                str(idx + 1),
                _format_date(d["date"]),
                day_name,
                str(d["light"]),
                str(d["heavy"]),
                str(tot),
                bar,
            ])

        _add_styled_table(doc, daily_headers, daily_rows,
                          col_widths=[0.5, 1.8, 2.2, 1.3, 1.3, 1.3, 3.5],
                          right_align_cols=[3, 4, 5],
                          bold_col=5,
                          bar_col=6)
        _add_spacer(doc, 10)

    # ── Poyezd harakati ──────────────────────────────────────
    _block_label(doc, t("rpt.train_movement").upper(), HDR_TEAL)

    def _fmt_dur(secs):
        if not secs:
            return "—"
        m = int(secs) // 60
        s = int(secs) % 60
        return (f"{m} {t('unit.min')} {s:02d} {t('unit.sec')}"
                if m > 0 else f"{s} {t('unit.sec')}")

    # 4 ta poyezd kartasi
    tcards = doc.add_table(rows=1, cols=4)
    tcards.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(tcards)

    _add_stat_card(tcards.rows[0].cells[0], str(train_stats["count"]),
                   t("rpt.total_trains").upper(), LIGHT_TEAL_BG,    ACCENT_TEAL,   ACCENT_TEAL_HEX)
    _add_stat_card(tcards.rows[0].cells[1],
                   _fmt_dur(train_stats["avg"]) if train_stats["count"] else "—",
                   t("rpt.avg_time").upper(), LIGHT_BLUE_BG,   BRAND_LIGHT,   ACCENT_BLUE_HEX)
    _add_stat_card(tcards.rows[0].cells[2],
                   _fmt_dur(train_stats["min"]) if train_stats["count"] else "—",
                   t("rpt.min_time").upper(),  LIGHT_GREEN_BG,  ACCENT_GREEN,  ACCENT_GREEN_HEX)
    _add_stat_card(tcards.rows[0].cells[3],
                   _fmt_dur(train_stats["max"]) if train_stats["count"] else "—",
                   t("rpt.max_time").upper(), LIGHT_ORANGE_BG, ACCENT_ORANGE, ACCENT_ORANGE_HEX)

    _add_spacer(doc, 8)

    # Poyezdlar o'tish jadvali
    train_events = stats_db.get_train_events_range(cid, date_from, date_to)
    if train_events:
        t_headers = [t("rpt.col_num"), t("rpt.col_date"), t("rpt.col_enter"),
                     t("rpt.col_exit"), t("rpt.col_duration")]
        t_rows = []
        for idx, ev in enumerate(train_events):
            t_rows.append([
                str(idx + 1),
                ev["date"],
                ev["start"],
                ev["end"],
                ev["duration_fmt"],
            ])
        _add_styled_table(doc, t_headers, t_rows,
                          col_widths=[0.6, 2.4, 2.0, 2.0, 2.8],
                          right_align_cols=[2, 3, 4],
                          bold_col=4)
    elif train_stats["count"] == 0:
        p = doc.add_paragraph()
        run = p.add_run("  " + t("rpt.no_train_period"))
        run.font.size = Pt(9)
        run.font.color.rgb = TEXT_GRAY
        run.italic = True
        _set_spacing(p, before=0, after=4)


# ═══════════════════════════════════════════════════════════
#  FOOTER
# ═══════════════════════════════════════════════════════════

def _add_footer(doc):
    _add_spacer(doc, 12)

    # Footer chiziq
    _add_colored_line(doc, GRAY_LINE, width=1)
    _add_spacer(doc, 4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— " + t("rpt.footer_end") + " —")
    run.font.size = Pt(10)
    run.font.color.rgb = TEXT_LIGHT
    run.italic = True
    _set_spacing(p, before=0, after=4)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run(t("rpt.footer_system") + "  |  ")
    run.font.size = Pt(8)
    run.font.color.rgb = TEXT_LIGHT
    run = p2.add_run(t("rpt.report_created_at", dt=datetime.now().strftime('%d.%m.%Y %H:%M')))
    run.font.size = Pt(8)
    run.font.color.rgb = TEXT_LIGHT


# ═══════════════════════════════════════════════════════════
#  STYLED TABLE
# ═══════════════════════════════════════════════════════════

def _add_styled_table(doc, headers, rows_data, col_widths=None,
                      right_align_cols=None, bold_col=None,
                      header_color=None, bar_col=None):
    """Professional jadval: rangli header, zebra qatorlar, vizual bar."""
    right_align_cols = right_align_cols or []
    if header_color is None:
        header_color = TABLE_HDR_BG   # HTML: #f1f5f9 light gray
    n_cols = len(headers)
    n_rows = len(rows_data) + 1

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            if i < n_cols:
                for row in table.rows:
                    row.cells[i].width = Cm(w)

    # ─── Header ─── HTML: thead tr { background:#f1f5f9 } yoki rangli
    use_light_hdr = (header_color == TABLE_HDR_BG)
    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_shading(cell, header_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(p, before=3, after=3)

        run = p.add_run(h_text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = TEXT_MID if use_light_hdr else WHITE
        run.font.name = 'Calibri'

    # ─── Data qatorlari ───
    for r_idx, row_data in enumerate(rows_data):
        bg = ROW_EVEN if r_idx % 2 == 0 else ROW_ODD
        for c_idx in range(n_cols):
            cell = table.rows[r_idx + 1].cells[c_idx]
            _set_cell_shading(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p = cell.paragraphs[0]
            if c_idx in right_align_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_spacing(p, before=2, after=2)

            value = row_data[c_idx] if c_idx < len(row_data) else ""
            run = p.add_run(value)
            run.font.name = 'Calibri'

            if c_idx == bar_col:
                # Bar grafik — teal rang
                run.font.size = Pt(7)
                run.font.color.rgb = ACCENT_TEAL
            else:
                run.font.size = Pt(9)
                run.font.color.rgb = TEXT_DARK

            if bold_col is not None and c_idx == bold_col:
                run.bold = True
                run.font.color.rgb = BRAND

    _set_table_borders(table)


# ═══════════════════════════════════════════════════════════
#  HELPERS — Components
# ═══════════════════════════════════════════════════════════

def _section_header(doc, num, text):
    """HTML .sec-title uslubi — chapda ko'k chiziq + sarlavha matni."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_table_borders(tbl)
    _set_table_width(tbl, Cm(17))

    # Chap: ko'k vertikal chiziq (HTML ::before — 4px blue bar)
    c0 = tbl.rows[0].cells[0]
    _set_cell_shading(c0, ACCENT_BLUE_HEX)
    c0.width = Cm(0.15)
    _set_cell_padding(c0, top=80, bottom=80, left=0, right=0)
    p0 = c0.paragraphs[0]
    p0.add_run("")
    _set_spacing(p0, before=0, after=0)

    # O'ng: sarlavha
    c1 = tbl.rows[0].cells[1]
    _set_cell_shading(c1, "FFFFFF")
    _set_cell_padding(c1, top=80, bottom=80, left=180, right=200)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p1.add_run(f"{num}.  {text}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = BRAND        # HTML: #1e3a5f
    run.font.name = 'Calibri'
    _set_spacing(p1, before=0, after=0)

    _add_spacer(doc, 8)


def _sub_header(doc, text):
    """Kichik sarlavha — chapda rangli chiziq."""
    p = doc.add_paragraph()
    _set_spacing(p, before=4, after=4)
    run = p.add_run("▌ ")
    run.font.color.rgb = BRAND_LIGHT
    run.font.size = Pt(11)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_DARK
    run.font.name = 'Calibri'


def _block_label(doc, text, color_hex):
    """HTML .sub-title uslubi — qalin matn + chap rangli chegara."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_table_borders(tbl)
    _set_table_width(tbl, Cm(17))

    # Chap chiziq
    c0 = tbl.rows[0].cells[0]
    _set_cell_shading(c0, color_hex)
    c0.width = Cm(0.1)
    _set_cell_padding(c0, top=60, bottom=60, left=0, right=0)
    c0.paragraphs[0].add_run("")
    _set_spacing(c0.paragraphs[0], before=0, after=0)

    # Matn
    c1 = tbl.rows[0].cells[1]
    _set_cell_shading(c1, "FFFFFF")
    _set_cell_padding(c1, top=60, bottom=60, left=160, right=160)
    p = c1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(p, before=0, after=0)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_MID     # HTML: .sub-title #374151
    run.font.name = 'Calibri'

    _add_spacer(doc, 4)


def _add_banner_card(cell, value, label, accent_hex):
    """Title sahifasidagi info karta — HTML: header-bottom-bar hb-item uslubi."""
    _set_cell_shading(cell, HDR_BG)
    _set_cell_padding(cell, top=160, bottom=160, left=120, right=120)
    _set_cell_top_border(cell, accent_hex, sz=14)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(value)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = WHITE
    run.font.name = 'Calibri'
    _set_spacing(p, before=0, after=2)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run(label.upper())
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x6A, 0x94, 0xB8)
    run.font.name = 'Calibri'
    _set_spacing(p2, before=0, after=0)


def _add_stat_card(cell, value, label, bg_color, value_color,
                   accent_hex=None):
    """Statistika karta — HTML .card uslubi: oq fon, qalin rangli tepa chegara."""
    _set_cell_shading(cell, CARD_BG)          # HTML: white
    _set_cell_padding(cell, top=140, bottom=140, left=120, right=120)

    if accent_hex:
        _set_cell_top_border(cell, accent_hex, sz=18)  # HTML: border-top: 4px

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(p, before=0, after=2)
    run = p.add_run(label)
    run.font.size = Pt(8)
    run.font.color.rgb = TEXT_GRAY                     # HTML: .card-lbl #6b7280
    run.font.name = 'Calibri'
    run.bold = True

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(p2, before=0, after=0)
    run = p2.add_run(value)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = TEXT_DARK                     # HTML: .card-val #111827
    run.font.name = 'Calibri'


# ═══════════════════════════════════════════════════════════
#  HELPERS — Low level
# ═══════════════════════════════════════════════════════════

def _format_date(date_str):
    """2026-02-01 → 01.02.2026"""
    try:
        d = datetime.strptime(date_str, "%Y-%m-%d")
        return d.strftime("%d.%m.%Y")
    except (ValueError, TypeError):
        return date_str


def _set_cell_shading(cell, color_hex):
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_top_border(cell, color_hex, sz=12):
    """Katakning faqat tepa chegarasini qo'yish (accent stripe)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>'
        f'</w:tcBorders>')
    tc_pr.append(borders)


def _set_cell_padding(cell, top=0, bottom=0, left=0, right=0):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>')
    tc_pr.append(margins)


def _set_table_width(table, width):
    """Jadval umumiy kengligini o'rnatish."""
    try:
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
            tbl.append(tbl_pr)
        tbl_w = parse_xml(
            f'<w:tblW {nsdecls("w")} w:w="{int(width)}" w:type="dxa"/>')
        tbl_pr.append(tbl_w)
    except Exception:
        pass


def _set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top    w:val="single" w:sz="4" w:space="0" w:color="{GRAY_LINE}"/>'
        f'  <w:left   w:val="single" w:sz="4" w:space="0" w:color="{GRAY_LINE}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{GRAY_LINE}"/>'
        f'  <w:right  w:val="single" w:sz="4" w:space="0" w:color="{GRAY_LINE}"/>'
        f'  <w:insideH w:val="single" w:sz="2" w:space="0" w:color="{GRAY_LINE}"/>'
        f'  <w:insideV w:val="single" w:sz="2" w:space="0" w:color="{GRAY_LINE}"/>'
        f'</w:tblBorders>')
    tbl_pr.append(borders)


def _remove_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top    w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left   w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right  w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>')
    tbl_pr.append(borders)


def _add_colored_line(doc, color_hex, width=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, before=0, after=0)
    pPr = p._p.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{width * 4}" w:space="1" w:color="{color_hex}"/>'
        f'</w:pBdr>')
    pPr.append(borders)


def _add_spacer(doc, pt_size):
    p = doc.add_paragraph("")
    _set_spacing(p, before=0, after=pt_size)


def _set_spacing(paragraph, before=0, after=0):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
