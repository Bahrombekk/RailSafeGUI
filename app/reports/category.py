# -*- coding: utf-8 -*-
"""
Temir yo'l kesishmasi TOIFASINI aniqlash (O'zbekiston standarti, 1-ILOVA).

Toifa (I-IV) sutkalik poyezd soni × keltirilgan transport soni bo'yicha
standart jadvaldan topiladi:

  - 1-jadval: umumiy foydalanishdagi kesishmalar (public)
  - 2-jadval: umumiy foydalanishda bo'lmagan kesishmalar (private)
  - 3-jadval: transport turlari koeffitsientlari (keltirilgan intensivlik)

Eslatma: detektor faqat YENGIL/OG'IR ni ajratadi (tonnajni bilmaydi), shuning
uchun keltirilgan transport = yengil*1.0 + og'ir*K (K sozlanadi, default 2.5).
"""

from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

logger = logging.getLogger("RailSafe.category")

INF = float("inf")

# Yengil koeffitsienti (3-jadval): yengil avtomobil = 1.0
LIGHT_COEF = 1.0
# Og'ir transport uchun default koeffitsient (yuk mashina o'rtacha ~2.5).
# Config: settings.toifa_heavy_coef bilan o'zgartiriladi.
DEFAULT_HEAVY_COEF = 2.5

# ── 1-jadval: umumiy foydalanish (public) ─────────────────────────────
# Poyezd qatori chegaralari (yuqori chegara, yorliq)
_PUBLIC_TRAINS = [(16, "16 tagacha"), (100, "17—100"),
                  (200, "101—200"), (INF, "200 dan ortiq")]
# Transport ustuni chegaralari
_PUBLIC_VEH = [(200, "200 tagacha"), (1000, "201—1000"), (3000, "1001—3000"),
               (7000, "3001—7000"), (INF, "7000 dan ortiq")]
# matritsa[poyezd_qatori][transport_ustuni] = toifa
_PUBLIC_MATRIX = [
    ["IV", "IV", "IV", "III", "II"],   # 16 tagacha
    ["IV", "IV", "III", "II", "I"],    # 17—100
    ["IV", "III", "II", "I", "I"],     # 101—200
    ["III", "II", "I", "I", "I"],      # 200 dan ortiq
]

# ── 2-jadval: umumiy foydalanishda emas (private) ─────────────────────
_PRIVATE_TRAINS = [(8, "8 tagacha"), (24, "9—24"),
                   (38, "25—38"), (INF, "38 dan ortiq")]
_PRIVATE_VEH = [(100, "100 gacha"), (500, "101—500"),
                (1000, "501—1000"), (INF, "1001 dan ortiq")]
_PRIVATE_MATRIX = [
    ["IV", "IV", "IV", "III"],   # 8 tagacha
    ["IV", "IV", "III", "II"],   # 9—24
    ["IV", "III", "II", "I"],    # 25—38
    ["III", "II", "I", "I"],     # 38 dan ortiq
]


def _band(value: float, bands: list) -> tuple[int, str]:
    """value qaysi diapazonga tushishini topadi. Returns (indeks, yorliq)."""
    for i, (upper, label) in enumerate(bands):
        if value <= upper:
            return i, label
    return len(bands) - 1, bands[-1][1]


def classify(avg_trains: float, avg_vehicles: float, public: bool = True):
    """O'rtacha sutkalik poyezd va (keltirilgan) transport bo'yicha toifani
    aniqlaydi.
    Returns dict: {toifa, train_band, veh_band}."""
    if public:
        trains, veh, matrix = _PUBLIC_TRAINS, _PUBLIC_VEH, _PUBLIC_MATRIX
    else:
        trains, veh, matrix = _PRIVATE_TRAINS, _PRIVATE_VEH, _PRIVATE_MATRIX
    ti, tlabel = _band(avg_trains, trains)
    vi, vlabel = _band(avg_vehicles, veh)
    return {"toifa": matrix[ti][vi], "train_band": tlabel, "veh_band": vlabel}


def compute_category(stats_db, crossing_id: int, date_from: str, date_to: str,
                     heavy_coef: float = DEFAULT_HEAVY_COEF,
                     public: bool = True) -> dict:
    """Bitta kesishma uchun toifani hisoblaydi.
    Returns dict: days, total_light, total_heavy, avg_trains, avg_reduced,
    toifa, train_band, veh_band."""
    daily = stats_db.get_date_range_daily(crossing_id, date_from, date_to)
    days = len(daily) if daily else 0
    total_light = sum(d["light"] for d in daily)
    total_heavy = sum(d["heavy"] for d in daily)
    train_stats = stats_db.get_train_range_stats(crossing_id, date_from, date_to)
    total_trains = train_stats.get("count", 0)

    if days <= 0:
        return {"days": 0, "total_light": total_light, "total_heavy": total_heavy,
                "avg_trains": 0.0, "avg_reduced": 0.0,
                "toifa": "—", "train_band": "—", "veh_band": "—"}

    reduced_total = total_light * LIGHT_COEF + total_heavy * heavy_coef
    avg_trains = total_trains / days
    avg_reduced = reduced_total / days
    result = classify(avg_trains, avg_reduced, public=public)
    return {
        "days": days,
        "total_light": total_light,
        "total_heavy": total_heavy,
        "avg_trains": avg_trains,
        "avg_reduced": avg_reduced,
        "toifa": result["toifa"],
        "train_band": result["train_band"],
        "veh_band": result["veh_band"],
    }


# ── DOCX hisobot ──────────────────────────────────────────────────────

def _fmt(n: float) -> str:
    """Sonni bo'sh joy bilan ajratib formatlash: 40520 -> '40 520'."""
    return f"{int(round(n)):,}".replace(",", " ")


def generate_category_report(config_manager, stats_db, date_from: str,
                             date_to: str, output_path: str,
                             heavy_coef: Optional[float] = None,
                             public: bool = True) -> bool:
    """Barcha kesishmalar uchun TOIFA hisobotini .docx qilib yaratadi."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if heavy_coef is None:
            heavy_coef = float(config_manager.get_settings().get(
                "toifa_heavy_coef", DEFAULT_HEAVY_COEF))

        doc = Document()

        title = doc.add_heading("Temir yo'l kesishmalarining toifalari", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = doc.add_paragraph(
            f"Hisobot davri: {_fmt_date(date_from)} — {_fmt_date(date_to)}  |  "
            f"Og'ir transport koeffitsienti (K): {heavy_coef}")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

        crossings = config_manager.get_crossings()
        for cr in crossings:
            cid = cr.get("id")
            name = cr.get("name", f"Pereezd_{cid}")
            cat = compute_category(stats_db, cid, date_from, date_to,
                                   heavy_coef=heavy_coef, public=public)

            h = doc.add_heading(f"{name} — {cat['toifa']} toifa", level=1)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            doc.add_paragraph(f"{name} — {cat['toifa']} toifa kesishma.")

            table = doc.add_table(rows=0, cols=2)
            table.style = "Light Grid Accent 1"
            rows = [
                ("Ko'rsatkich", "Qiymat"),
                ("To'liq kunlar (24h)", f"{cat['days']} kun"),
                ("O'rtacha transport/sutka", f"{_fmt(cat['avg_reduced'])} (koeffitsient bilan)"),
                ("O'rtacha poyezd/sutka", f"{int(round(cat['avg_trains']))} ta"),
                ("Poyezd qatori", cat["train_band"]),
                ("Transport ustuni", cat["veh_band"]),
                ("TOIFA", cat["toifa"]),
            ]
            for k, v in rows:
                cells = table.add_row().cells
                cells[0].text = k
                cells[1].text = str(v)
            # sarlavha qatorini qalin qilish
            for cell in table.rows[0].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
            doc.add_paragraph("")

        note = doc.add_paragraph(
            "Izoh: keltirilgan transport = yengil×1.0 + og'ir×K. Detektor tonnajni "
            "ajratmagani uchun og'ir transportga o'rtacha koeffitsient (K) qo'llaniladi.")
        for r in note.runs:
            r.font.size = Pt(8)

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return True
    except Exception as e:
        logger.exception("Toifa hisoboti xato: %s", e)
        return False


def _fmt_date(s: str) -> str:
    try:
        return datetime.strptime(s[:10], "%Y-%m-%d").strftime("%d.%m.%Y")
    except Exception:
        return s
